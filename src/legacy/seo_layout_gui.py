import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from datetime import datetime
import os
import re
import json
from docx import Document
from jinja2 import Template
try:
    import ttkbootstrap as tb
    from ttkbootstrap.widgets import DateEntry as TBDateEntry
    HAS_TTKBOOTSTRAP = True
except Exception:
    TBDateEntry = None
    HAS_TTKBOOTSTRAP = False
try:
    import webview
    HAS_WEBVIEW = True
except Exception:
    HAS_WEBVIEW = False

try:
    from tkcalendar import DateEntry as TKCalDateEntry
    HAS_TKCALENDAR = True
except Exception:
    TKCalDateEntry = None
    HAS_TKCALENDAR = False

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
ARTICLE_NUMBER_FILE = os.path.join(BASE_DIR, "article_number.txt")
SETTINGS_FILE = os.path.join(BASE_DIR, "settings.json")

def get_article_number():
    if not os.path.exists(ARTICLE_NUMBER_FILE):
        with open(ARTICLE_NUMBER_FILE, "w", encoding="utf-8") as f:
            f.write("1")
        return 1
    with open(ARTICLE_NUMBER_FILE, "r", encoding="utf-8") as f:
        try:
            num = int(f.read().strip())
        except Exception:
            num = 1
    return num

def set_article_number(num):
    with open(ARTICLE_NUMBER_FILE, "w", encoding="utf-8") as f:
        f.write(str(num))

# tp 標記正則
TP_H1 = r'\(tp_h1\)(.*)'
TP_H2 = r'\(tp_h2\)(.*)'
TP_H3 = r'\(tp_h3\)(.*)'
TP_SEC = r'\(tp_sec\)'
TP_SEC_QA = r'\(tp_sec_qa\)'
TP_H3_Q = r'\(tp_h3_q\)(.*)'
TP_ANS = r'\(tp_ans\)(.*)'

def parse_docx(docx_path):
    doc = Document(docx_path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    headline = ""
    description = ""
    faq_list = []
    in_faq = False
    current_q = None
    current_a = None

    for line in lines:
        # 主標題
        if re.match(TP_H1, line):
            headline = re.match(TP_H1, line).group(1).strip()
        # 描述（第一個 tp_h2）
        elif re.match(TP_H2, line) and not description:
            description = re.match(TP_H2, line).group(1).strip()
        # FAQ 區塊開始
        elif re.match(TP_SEC_QA, line):
            in_faq = True
        # FAQ 問題
        elif in_faq and re.match(TP_H3_Q, line):
            current_q = re.match(TP_H3_Q, line).group(1).strip()
        # FAQ 答案
        elif in_faq and re.match(TP_ANS, line):
            current_a = re.match(TP_ANS, line).group(1).strip()
            if current_q and current_a:
                faq_list.append({"question": current_q, "answer": current_a})
                current_q = None
                current_a = None
        # FAQ 區塊結束
        elif in_faq and re.match(TP_SEC, line):
            in_faq = False

    # 若 description 沒抓到，取 headline 前 100 字
    if not description and headline:
        description = headline[:100]
    return headline, description, faq_list

def update_jsonld(template_html, context, faq_list):
    # 抽取 <script type="application/ld+json">
    script_re = r'<script type="application/ld\+json">\s*(\{.*?\})\s*</script>'
    match = re.search(script_re, template_html, re.DOTALL)
    if not match:
        return template_html  # 沒有 JSON-LD 區塊
    jsonld = json.loads(match.group(1))
    # 更新欄位
    jsonld['author']['name'] = context['author_name']
    jsonld['dateModified'] = context['TheDate']
    jsonld['datePublished'] = context['TheDate']
    jsonld['headline'] = context['headline']
    jsonld['description'] = context['description']
    jsonld['mainEntityOfPage']['@id'] = f"https://pm.shiny.com.tw/news-detail.php?id={context['article_number']}"
    jsonld['publisher']['name'] = context['OrganizationName']
    # FAQ mainEntity
    faq_entities = []
    for faq in faq_list:
        faq_entities.append({
            "@type": "Question",
            "name": faq["question"],
            "acceptedAnswer": {
                "@type": "Answer",
                "text": faq["answer"]
            }
        })
    jsonld['mainEntity']['mainEntity'] = faq_entities
    # 回寫
    new_jsonld = json.dumps(jsonld, ensure_ascii=False, indent=4)
    new_html = re.sub(script_re, f'<script type="application/ld+json">{new_jsonld}</script>', template_html, flags=re.DOTALL)
    return new_html

def insert_faq_section(html, faq_list):
    # 若沒有 <section id="faq-section">，插入 FAQ 區塊
    if '<section id="faq-section">' not in html:
        faq_html = '<section id="faq-section">\n<h2>常見問答 (Q&amp;A)</h2>\n'
        for faq in faq_list:
            faq_html += f'<h3>{faq["question"]}</h3>\n<p>{faq["answer"]}</p>\n'
        faq_html += '</section>\n'
        html = html.replace('</article>', f'{faq_html}</article>')
    return html

class SEOLayoutGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("SEO 文章模板 GUI 工具")
        if HAS_TTKBOOTSTRAP:
            try:
                tb.Style(theme='flatly')
            except Exception:
                pass

        # 主題切換（僅在 ttkbootstrap 可用時顯示）
        if HAS_TTKBOOTSTRAP:
            theme_bar = ttk.Frame(root)
            theme_bar.pack(fill="x", padx=8, pady=4)
            ttk.Label(theme_bar, text="主題：").pack(side="left")
            self.theme_var = tk.StringVar(value='flatly')
            theme_combo = ttk.Combobox(theme_bar, textvariable=self.theme_var, width=12, state='readonly')
            theme_combo['values'] = ('flatly', 'cosmo', 'darkly', 'morph', 'sandstone', 'solar')
            theme_combo.pack(side="left")
            theme_combo.bind('<<ComboboxSelected>>', self._on_theme_change)
        # Notebook 分頁
        self.notebook = ttk.Notebook(root)
        self.tab_basic = ttk.Frame(self.notebook)
        self.tab_preview = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_basic, text="基本與檔案")
        self.notebook.add(self.tab_preview, text="FAQ 預覽")
        self.notebook.pack(fill="both", expand=True, padx=8, pady=8)

        # 基本資訊分頁
        ttk.Label(self.tab_basic, text="作者").grid(row=0, column=0, sticky="e")
        self.author = ttk.Entry(self.tab_basic, width=30)
        self.author.insert(0, "炫麗鑫")
        self.author.grid(row=0, column=1, sticky="w")

        ttk.Label(self.tab_basic, text="文章日期 (YYYY-MM-DD)").grid(row=0, column=2, sticky="e")
        today = datetime.today().strftime('%Y-%m-%d')
        if TBDateEntry:
            self.date = TBDateEntry(self.tab_basic, bootstyle="info", dateformat='%Y-%m-%d')
            self.date.entry.delete(0, tk.END)
            self.date.entry.insert(0, today)
        elif HAS_TKCALENDAR and TKCalDateEntry:
            self.date = TKCalDateEntry(self.tab_basic, date_pattern='yyyy-mm-dd')
            self.date.set_date(today)
        else:
            self.date = ttk.Entry(self.tab_basic, width=20)
            self.date.insert(0, today)
        self.date.grid(row=0, column=3, sticky="w")

        ttk.Label(self.tab_basic, text="標題 (headline)").grid(row=1, column=0, sticky="e")
        self.headline = ttk.Entry(self.tab_basic, width=60)
        self.headline.grid(row=1, column=1, columnspan=3, sticky="we")

        ttk.Label(self.tab_basic, text="描述 (description)").grid(row=2, column=0, sticky="e")
        self.description = ttk.Entry(self.tab_basic, width=60)
        self.description.grid(row=2, column=1, columnspan=3, sticky="we")

        ttk.Label(self.tab_basic, text="組織名稱").grid(row=3, column=0, sticky="e")
        self.organization = ttk.Entry(self.tab_basic, width=30)
        self.organization.insert(0, "Shiny黃金白銀")
        self.organization.grid(row=3, column=1, sticky="w")

        ttk.Label(self.tab_basic, text="文章編號").grid(row=3, column=2, sticky="e")
        self.article_number = ttk.Entry(self.tab_basic, width=20)
        self.article_number.grid(row=3, column=3, sticky="w")
        current = get_article_number() + 1
        self.article_number.insert(0, str(current))
        ttk.Button(self.tab_basic, text="確認號碼", command=self.confirm_article_number).grid(row=3, column=4, padx=4)

        # 檔案與輸出（合併在同頁）
        ttk.Label(self.tab_basic, text="模板檔案").grid(row=4, column=0, sticky="e")
        self.template_path = ttk.Entry(self.tab_basic, width=60)
        default_template = os.path.join(BASE_DIR, "templates", "seo_layout.html")
        self.template_path.insert(0, default_template if os.path.exists(default_template) else "")
        self.template_path.grid(row=4, column=1, columnspan=2, sticky="we")
        ttk.Button(self.tab_basic, text="瀏覽", command=self.select_template).grid(row=4, column=3, sticky="w")

        ttk.Label(self.tab_basic, text="Word 檔案").grid(row=5, column=0, sticky="e")
        self.word_path = ttk.Entry(self.tab_basic, width=60)
        self.word_path.grid(row=5, column=1, columnspan=2, sticky="we")
        ttk.Button(self.tab_basic, text="瀏覽", command=self.select_word).grid(row=5, column=3, sticky="w")

        # 動作按鈕
        self.parse_btn = ttk.Button(self.tab_basic, text="解析 FAQ 預覽", command=self.update_preview)
        self.parse_btn.grid(row=6, column=0, pady=6)
        self.preview_btn = ttk.Button(self.tab_basic, text="預覽 HTML（瀏覽器）", command=self.preview_in_browser)
        self.preview_btn.grid(row=6, column=1, pady=6)
        self.generate_btn = ttk.Button(self.tab_basic, text="產生 HTML", command=self.generate_html)
        self.generate_btn.grid(row=6, column=2, pady=6)

        # FAQ 預覽分頁
        self.preview_text = tk.Text(self.tab_preview, wrap="word", height=18)
        self.preview_text.pack(fill="both", expand=True)
        if HAS_WEBVIEW:
            self.webview_btn = ttk.Button(self.tab_preview, text="內嵌預覽（WebView）", command=self.preview_in_webview)
            self.webview_btn.pack(anchor="w", pady=6)

        # 載入設定並初始化驗證
        self._load_settings()
        date_widget = self.date.entry if TBDateEntry and isinstance(self.date, TBDateEntry) else self.date
        for widget in [self.author, date_widget, self.headline, self.description, self.organization, self.article_number, self.template_path, self.word_path]:
            widget.bind('<KeyRelease>', lambda e: self._update_generate_state())
        self._update_generate_state()

    def select_template(self):
        path = filedialog.askopenfilename(filetypes=[("HTML Files", "*.html")], initialdir=os.path.join(BASE_DIR, "templates"))
        if path:
            self.template_path.delete(0, tk.END)
            self.template_path.insert(0, path)
            self._save_settings()

    def select_word(self):
        path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")], initialdir=os.path.join(BASE_DIR, "input_docs"))
        if path:
            self.word_path.delete(0, tk.END)
            self.word_path.insert(0, path)
            # 自動解析主標題與描述
            headline, description = parse_main_content_from_word(path)
            if headline:
                self.headline.delete(0, tk.END)
                self.headline.insert(0, headline)
            if description:
                self.description.delete(0, tk.END)
                self.description.insert(0, description)
            self._save_settings()

    def confirm_article_number(self):
        value = self.article_number.get().strip()
        if not value.isdigit():
            messagebox.showerror("錯誤", "文章編號需為數字")
            return
        set_article_number(int(value))
        messagebox.showinfo("完成", f"已更新文章編號至 {value}")
        self._save_settings()

    def _parse_faq_from_word(self, word_path):
        doc = Document(word_path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]

        faq_questions = []
        faq_answers = []
        in_qa = False
        current_q = None
        collecting_ans = False
        ans_lines = []

        for line in lines:
            if line.startswith('(tp_sec_qa)'):
                in_qa = True
                current_q = None
                collecting_ans = False
                ans_lines = []
                continue
            if not in_qa:
                continue
            if line.startswith('(tp_h3_q)'):
                if current_q is not None and ans_lines:
                    faq_questions.append(current_q)
                    faq_answers.append('\n'.join(ans_lines).strip())
                current_q = line.replace('(tp_h3_q)', '').strip()
                ans_lines = []
                collecting_ans = False
                continue
            if line.startswith('(tp_ans)'):
                collecting_ans = True
                ans_lines = []
                continue
            if re.match(r'\(tp_[a-zA-Z0-9_]+\)', line):
                if current_q is not None and ans_lines:
                    faq_questions.append(current_q)
                    faq_answers.append('\n'.join(ans_lines).strip())
                current_q = None
                ans_lines = []
                collecting_ans = False
                continue
            if collecting_ans:
                ans_lines.append(line.strip())

        if current_q is not None and ans_lines:
            faq_questions.append(current_q)
            faq_answers.append('\n'.join(ans_lines).strip())

        return faq_questions, faq_answers

    def update_preview(self):
        path = self.word_path.get().strip()
        if not path or not os.path.exists(path):
            messagebox.showerror("錯誤", "請先選擇 Word 檔")
            return
        questions, answers = self._parse_faq_from_word(path)
        self.preview_text.delete('1.0', tk.END)
        if not questions:
            self.preview_text.insert(tk.END, "未解析到 FAQ 內容。\n請確認 Word 中的 (tp_sec_qa)/(tp_h3_q)/(tp_ans) 標記。")
        else:
            for i, (q, a) in enumerate(zip(questions, answers), start=1):
                self.preview_text.insert(tk.END, f"Q{i}: {q}\n")
                self.preview_text.insert(tk.END, f"A{i}: {a}\n\n")

    def _replace_faq_in_html(self, html, questions, answers):
        # HTML FAQ section 內容替換（以 id="faq-section" 區塊為準）；若無則新增於 </article> 前
        section_pattern = re.compile(r'(<section[^>]*id="faq-section"[^>]*>.*?</section>)', re.DOTALL)
        m = section_pattern.search(html)
        faq_html = ""
        for q, a in zip(questions, answers):
            faq_html += f'<h3>{q}</h3>\n<p>{a}</p>\n'
        if m:
            section_html = m.group(1)
            section_new = re.sub(r'(<h3>.*?</h3>\s*<p>.*?</p>\s*)+', faq_html, section_html, flags=re.DOTALL)
            html = html.replace(section_html, section_new)
        else:
            block = '<section id="faq-section">\n<h2>常見問答 (Q&amp;A)</h2>\n' + faq_html + '</section>\n'
            if '</article>' in html:
                html = html.replace('</article>', block + '</article>')
            else:
                html += "\n" + block
        return html

    def _update_jsonld(self, html, author, the_date, headline, description, organization, article_number, questions, answers):
        # 尋找第一個 JSON-LD script，解析並安全更新
        script_pattern = re.compile(r'(<script[^>]*type=\"application/ld\+json\"[^>]*>)([\s\S]*?)(</script>)', re.IGNORECASE)
        m = script_pattern.search(html)
        if not m:
            return html
        prefix, json_text, suffix = m.group(1), m.group(2), m.group(3)
        try:
            data = json.loads(json_text)
        except Exception:
            return html
        data["author"] = {"@type": "Organization", "name": author}
        data["dateModified"] = the_date
        data["datePublished"] = the_date
        data["headline"] = headline
        data["description"] = description
        data["mainEntityOfPage"] = {"@id": f"https://pm.shiny.com.tw/news-detail.php?id={article_number}", "@type": "WebPage"}
        data["publisher"] = {"@type": "Organization", "logo": {"@type": "ImageObject", "url": "https://pm.shiny.com.tw/images/logo.png"}, "name": organization}
        faq_items = []
        for q, a in zip(questions, answers):
            faq_items.append({
                "@type": "Question",
                "acceptedAnswer": {"@type": "Answer", "text": a if a else "無資料"},
                "name": q
            })
        data["mainEntity"] = {"@type": "FAQPage", "mainEntity": faq_items}
        new_json = json.dumps(data, ensure_ascii=False, indent=4)
        new_script = prefix + new_json + suffix
        return html[:m.start()] + new_script + html[m.end():]

    def generate_html(self):
        author = self.author.get().strip()
        the_date = self._get_date_text()
        organization = self.organization.get().strip()
        article_number = self.article_number.get().strip()
        template_path = self.template_path.get().strip()
        word_path = self.word_path.get().strip()

        if not template_path or not os.path.exists(template_path):
            messagebox.showerror("錯誤", "請選擇模板檔案")
            return
        if not word_path or not os.path.exists(word_path):
            messagebox.showerror("錯誤", "請選擇 Word 檔")
            return

        headline, description = parse_main_content_from_word(word_path)
        self.headline.delete(0, tk.END)
        self.headline.insert(0, headline)
        self.description.delete(0, tk.END)
        self.description.insert(0, description)

        questions, answers = self._parse_faq_from_word(word_path)
        intro_paragraphs = parse_intro_from_word(word_path)
        sections = parse_sections_from_word(word_path)

        # 產生各區塊內容
        with open(template_path, "r", encoding="utf-8") as f:
            html = f.read()

        # 1. h1
        h1_html = f"<h1>{headline}</h1>"

        # 2. intro-summary
        intro_html = '<section class="intro-summary">\n'
        if intro_paragraphs:
            has_intro_title, paragraphs = intro_paragraphs
            if has_intro_title:
                intro_html += '    <h2>前言</h2>\n'
            for p in paragraphs:
                intro_html += f'    <p>{p}</p>\n'
        intro_html += '</section>'

        # 3. 主文 section（數量不固定）
        sections_html = ""
        for sec in sections:
            sections_html += f"<section>\n{sec}\n</section>\n"

        # 4. FAQ
        faq_html = '<section id="faq-section">\n<h2>常見問答 (Q&amp;A)</h2>\n'
        for q, a in zip(questions, answers):
            faq_html += f'<h3>{q}</h3>\n<p>{a}</p>\n'
        faq_html += '</section>'

        # 5. footer（直接用樣板原本內容，不動）

        # 依照樣板順序分別覆蓋
        html = re.sub(r'<h1>.*?</h1>', h1_html, html, flags=re.DOTALL)
        html = re.sub(r'<section class="intro-summary">.*?</section>', intro_html, html, flags=re.DOTALL)

        # 只覆蓋主文 <section> 內容（不動 FAQ/footer/intro-summary）
        main_sections = re.findall(
            r'(<section>\s*<h2>.*?</h2>\s*</section>)', html, re.DOTALL
        )
        if main_sections:
            # 依照 sections 數量，逐一覆蓋 section 內容
            for i, old_sec in enumerate(main_sections):
                if i < len(sections):
                    new_sec = f"<section>\n{sections[i]}\n</section>"
                    html = html.replace(old_sec, new_sec, 1)
                else:
                    # 若 sections 不夠，剩下的 section 保留原本內容
                    break

        html = re.sub(r'<section id="faq-section">.*?</section>', faq_html, html, flags=re.DOTALL)

        # 其它變數替換
        html = html.replace("{{author_name}}", author)
        html = html.replace("{{TheDate}}", the_date)
        html = html.replace("{{headline}}", headline)
        html = html.replace("{{description}}", description)
        html = html.replace("{{OrganizationName}}", organization)
        html = html.replace("{{article_number}}", article_number)

        # FAQ JSON-LD 安全更新（如有）
        html = self._update_jsonld(html, author, the_date, headline, description, organization, article_number, questions, answers)

        # 預設輸出檔名
        today = datetime.today().strftime('%Y%m%d')
        base = f"output_{today}.html"
        candidate = base
        idx = 1
        out_dir = os.path.join(BASE_DIR, "output")
        os.makedirs(out_dir, exist_ok=True)
        while os.path.exists(os.path.join(out_dir, candidate)):
            candidate = f"output_{today}_{idx}.html"
            idx += 1

        out_path = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("HTML Files", "*.html")], initialfile=candidate, initialdir=out_dir)
        if out_path:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(html)
            messagebox.showinfo("完成", f"已產生 {out_path}")
            self._save_settings()

    def _on_theme_change(self, event=None):
        if not HAS_TTKBOOTSTRAP:
            return
        try:
            tb.Style().theme_use(self.theme_var.get())
        except Exception:
            pass

    def _update_generate_state(self):
        author = self.author.get().strip()
        the_date = self._get_date_text()
        headline = self.headline.get().strip()
        template_ok = os.path.exists(self.template_path.get().strip())
        word_ok = os.path.exists(self.word_path.get().strip())
        enable = bool(author and the_date and headline and template_ok and word_ok)
        state = tk.NORMAL if enable else tk.DISABLED
        try:
            self.generate_btn.config(state=state)
        except Exception:
            pass
        try:
            self.preview_btn.config(state=state)
        except Exception:
            pass

    def _load_settings(self):
        if not os.path.exists(SETTINGS_FILE):
            return
        try:
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for key, widget in {
                'author': self.author,
                'date': self.date,
                'headline': self.headline,
                'description': self.description,
                'organization': self.organization,
                'article_number': self.article_number,
                'template_path': self.template_path,
                'word_path': self.word_path,
            }.items():
                if key == 'date' and 'date' in data:
                    self._set_date_text(data.get('date', ''))
                elif key in data and isinstance(widget, (tk.Entry, ttk.Entry)):
                    widget.delete(0, tk.END)
                    widget.insert(0, data.get(key, ''))
        except Exception:
            pass

    def _save_settings(self):
        data = {
            'author': self.author.get().strip(),
            'date': self._get_date_text(),
            'headline': self.headline.get().strip(),
            'description': self.description.get().strip(),
            'organization': self.organization.get().strip(),
            'article_number': self.article_number.get().strip(),
            'template_path': self.template_path.get().strip(),
            'word_path': self.word_path.get().strip(),
        }
        try:
            with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def preview_in_browser(self):
        import webbrowser
        author = self.author.get().strip()
        the_date = self._get_date_text()
        organization = self.organization.get().strip()
        article_number = self.article_number.get().strip()
        template_path = self.template_path.get().strip()
        word_path = self.word_path.get().strip()

        if not template_path or not os.path.exists(template_path):
            messagebox.showerror("錯誤", "請選擇模板檔案")
            return
        if not word_path or not os.path.exists(word_path):
            messagebox.showerror("錯誤", "請選擇 Word 檔")
            return

        # 直接解析 Word 檔主標題與描述
        headline, description = parse_main_content_from_word(word_path)
        self.headline.delete(0, tk.END)
        self.headline.insert(0, headline)
        self.description.delete(0, tk.END)
        self.description.insert(0, description)

        questions, answers = self._parse_faq_from_word(word_path)
        intro_paragraphs = parse_intro_from_word(word_path)
        sections = parse_sections_from_word(word_path)

        # 產生各區塊內容
        with open(template_path, "r", encoding="utf-8") as f:
            html = f.read()

        # 1. h1
        h1_html = f"<h1>{headline}</h1>"

        # 2. intro-summary
        intro_html = '<section class="intro-summary">\n'
        if intro_paragraphs:
            has_intro_title, paragraphs = intro_paragraphs
            if has_intro_title:
                intro_html += '    <h2>前言</h2>\n'
            for p in paragraphs:
                intro_html += f'    <p>{p}</p>\n'
        intro_html += '</section>'

        # 3. 主文 section（數量不固定）
        sections_html = ""
        for sec in sections:
            sections_html += f"<section>\n{sec}\n</section>\n"

        # 4. FAQ
        faq_html = '<section id="faq-section">\n<h2>常見問答 (Q&amp;A)</h2>\n'
        for q, a in zip(questions, answers):
            faq_html += f'<h3>{q}</h3>\n<p>{a}</p>\n'
        faq_html += '</section>'

        # 依照樣板順序分別覆蓋
        html = re.sub(r'<h1>.*?</h1>', h1_html, html, flags=re.DOTALL)
        html = re.sub(r'<section class="intro-summary">.*?</section>', intro_html, html, flags=re.DOTALL)

        # 只覆蓋主文 <section> 內容（不動 FAQ/footer/intro-summary）
        main_sections = re.findall(
            r'(<section>\s*<h2>.*?</h2>\s*</section>)', html, re.DOTALL
        )
        if main_sections:
            # 依照 sections 數量，逐一覆蓋 section 內容
            for i, old_sec in enumerate(main_sections):
                if i < len(sections):
                    new_sec = f"<section>\n{sections[i]}\n</section>"
                    html = html.replace(old_sec, new_sec, 1)
                else:
                    # 若 sections 不夠，剩下的 section 保留原本內容
                    break

        html = re.sub(r'<section id="faq-section">.*?</section>', faq_html, html, flags=re.DOTALL)

        # 其它變數替換
        html = html.replace("{{author_name}}", author)
        html = html.replace("{{TheDate}}", the_date)
        html = html.replace("{{headline}}", headline)
        html = html.replace("{{description}}", description)
        html = html.replace("{{OrganizationName}}", organization)
        html = html.replace("{{article_number}}", article_number)

        # FAQ JSON-LD 安全更新（如有）
        html = self._update_jsonld(html, author, the_date, headline, description, organization, article_number, questions, answers)

        html = self._update_jsonld(html, author, the_date, headline, description, organization, article_number, questions, answers)

        # 儲存預覽檔案
        out_dir = os.path.join(BASE_DIR, "output")
        os.makedirs(out_dir, exist_ok=True)
        temp_path = os.path.join(out_dir, "preview_temp.html")
        with open(temp_path, "w", encoding="utf-8") as f:
            f.write(html)
        webbrowser.open_new_tab(temp_path)

    def preview_in_webview(self):
        if not HAS_WEBVIEW:
            messagebox.showwarning("提示", "未安裝 pywebview，已略過。")
            return
        author = self.author.get().strip()
        the_date = self._get_date_text()
        organization = self.organization.get().strip()
        article_number = self.article_number.get().strip()
        template_path = self.template_path.get().strip()
        word_path = self.word_path.get().strip()

        headline, description = parse_main_content_from_word(word_path)
        self.headline.delete(0, tk.END)
        self.headline.insert(0, headline)
        self.description.delete(0, tk.END)
        self.description.insert(0, description)

        questions, answers = self._parse_faq_from_word(word_path)
        sections = parse_sections_from_word(word_path)
        main_content_html = f"<h1>{headline}</h1>\n" + "\n".join([f"<section>\n{s}</section>" for s in sections])

        with open(template_path, "r", encoding="utf-8") as f:
            html = f.read()
        html = html.replace("{{author_name}}", author)
        html = html.replace("{{TheDate}}", the_date)
        html = html.replace("{{headline}}", headline)
        html = html.replace("{{description}}", description)
        html = html.replace("{{OrganizationName}}", organization)
        html = html.replace("{{article_number}}", article_number)

        html = re.sub(r'(<article[^>]*>)(.*?)(<section class="intro-summary">.*?</section>)(.*?)(<section id="faq-section">|<hr />|</article>)',
                      lambda m: f'{m.group(1)}\n{main_content_html}\n{m.group(3)}{m.group(5)}',
                      html, flags=re.DOTALL)

        html = self._update_jsonld(html, author, the_date, headline, description, organization, article_number, questions, answers)
        if questions and answers:
            html = self._replace_faq_in_html(html, questions, answers)
        window = webview.create_window('預覽', html=html)
        webview.start()

    def _get_date_text(self):
        # 支援 ttkbootstrap DateEntry / tkcalendar DateEntry / Entry
        if TBDateEntry and isinstance(self.date, TBDateEntry):
            try:
                return self.date.entry.get().strip()
            except Exception:
                return ''
        if HAS_TKCALENDAR and TKCalDateEntry and isinstance(self.date, TKCalDateEntry):
            try:
                return self.date.get().strip()
            except Exception:
                return ''
        try:
            return self.date.get().strip()
        except Exception:
            return ''

    def _set_date_text(self, value):
        try:
            if not value:
                value = datetime.today().strftime('%Y-%m-%d')
            if TBDateEntry and isinstance(self.date, TBDateEntry):
                self.date.entry.delete(0, tk.END)
                self.date.entry.insert(0, value)
            elif HAS_TKCALENDAR and TKCalDateEntry and isinstance(self.date, TKCalDateEntry):
                if isinstance(value, str):
                    value = datetime.strptime(value, '%Y-%m-%d')
                self.date.set_date(value)
            else:
                self.date.delete(0, tk.END)
                self.date.insert(0, value)
        except Exception:
            # 如果設置失敗，使用今天日期
            today = datetime.today().strftime('%Y-%m-%d')
            self._set_date_text(today)
        try:
            self.date.delete(0, tk.END)
            self.date.insert(0, value)
        except Exception:
            pass

def parse_main_content_from_word(word_path):
    doc = Document(word_path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    headline = ""
    description = ""
    for line in lines:
        if re.match(TP_H1, line):
            headline = re.match(TP_H1, line).group(1).strip()
        elif re.match(TP_H2, line) and not description:
            description = re.match(TP_H2, line).group(1).strip()
    # 若 description 沒抓到，取 headline 前 100 字
    if not description and headline:
        description = headline[:100]
    return headline, description

def parse_sections_from_word(word_path):
    doc = Document(word_path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    sections = []
    current_section = []  # 存儲當前section的內容列表
    in_section = False
    
    for i, line in enumerate(lines):
        # 遇到新的區塊標記
        if re.match(TP_SEC, line):
            # 保存之前的section（如果有的話）
            if current_section:
                section_html = "".join(current_section)
                sections.append(section_html)
                current_section = []
            in_section = True
            continue
            
        # 遇到其他區塊標記時結束當前section
        if in_section and (re.match(TP_SEC_QA, line) or re.match(TP_INTRO, line)):
            if current_section:
                section_html = "".join(current_section)
                sections.append(section_html)
                current_section = []
            in_section = False
            continue
            
        # 收集section中的內容
        if in_section:
            if re.match(TP_H2, line):
                h2_content = re.match(TP_H2, line).group(1).strip()
                current_section.append(f'<h2>{h2_content}</h2>\n')
            elif re.match(TP_H3, line):
                h3_content = re.match(TP_H3, line).group(1).strip()
                current_section.append(f'<h3>{h3_content}</h3>\n')
            elif not re.match(r'\(tp_.*\)', line):  # 一般段落內容
                current_section.append(f'<p>{line}</p>\n')
                
    # 保存最後一個section（如果有的話）
    if current_section:
        section_html = "".join(current_section)
        sections.append(section_html)
        
    return sections

def parse_intro_from_word(word_path):
    doc = Document(word_path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    intro_paragraphs = []
    in_intro = False
    has_intro_title = False
    
    for line in lines:
        if '(tp_intro)' in line:
            in_intro = True
            content = line.replace('(tp_intro)', '').strip()
            if content:
                intro_paragraphs.append(content)
        elif in_intro and '(tp_sec)' in line:
            # 遇到 (tp_sec) 就結束前言
            in_intro = False
        elif in_intro:
            # 在前言區塊中且還沒遇到 (tp_sec)，繼續收集內容
            if '(tp_h2)前言' in line:
                has_intro_title = True
            else:
                if line:
                    intro_paragraphs.append(line)
            
    return (has_intro_title, intro_paragraphs)

### tp_intro 標記解析

# - 支援 Word 檔中 (tp_intro) 標記，該段落會自動包進 `<section class="intro-summary">`，並插入於 `<h1>` 之後。

if __name__ == "__main__":
    try:
        if HAS_TTKBOOTSTRAP:
            root = tb.Window(themename='flatly')
        else:
            root = tk.Tk()
    except Exception:
        root = tk.Tk()
    app = SEOLayoutGUI(root)
    root.mainloop()
