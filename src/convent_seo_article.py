import docx
import os

def docx_to_html(docx_path, html_path):
    doc = docx.Document(docx_path)
    html = []

    # 僅生成內容主體，SEO 與樣式應由模板負責
    html.append('<article class="seo-article-content">')

    for para in doc.paragraphs:
        style = para.style.name.lower()
        text = para.text.strip()
        if not text:
            continue
        if "heading 1" in style:
            html.append(f"<h1>{text}</h1>")
        elif "heading 2" in style:
            html.append(f"<h2>{text}</h2>")
        elif "heading 3" in style:
            html.append(f"<h3>{text}</h3>")
        elif "list" in style:
            html.append(f"<ul><li>{text}</li></ul>")
        else:
            html.append(f"<p>{text}</p>")

    for table in doc.tables:
        html.append("<table>")
        for i, row in enumerate(table.rows):
            html.append("<tr>")
            for cell in row.cells:
                tag = "th" if i == 0 else "td"
                html.append(f"<{tag}>{cell.text.strip()}</{tag}>")
            html.append("</tr>")
        html.append("</table>")

    html.append("</article>")

    with open(html_path, "w", encoding="utf-8") as f:
        f.write("\n".join(html))

if __name__ == "__main__":
    # 範例：手動執行時才轉換
    sample_in = "0911見聞集錦_黃金的故事.docx"
    sample_out = "output.html"
    if os.path.exists(sample_in):
        docx_to_html(sample_in, sample_out)