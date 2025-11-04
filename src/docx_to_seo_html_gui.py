import tkinter as tk
from tkinter import filedialog, messagebox
import docx
import os
from datetime import datetime

# 產生 SEO 結構化資料
SEO_TEMPLATE = '''<script type="application/ld+json">
{{
    "@context": "https://schema.org",
    "@type": "Article",
    "author": {{"@type": "Organization", "name": "炫麗鑫"}},
    "dateModified": "{dateModified}",
    "datePublished": "{datePublished}",
    "description": "{description}",
    "headline": "{headline}",
    "image": "",
    "mainEntityOfPage": {{"@id": "https://pm.shiny.com.tw/news-detail.php?id={mainEntityId}", "@type": "WebPage"}},
    "publisher": {{"@type": "Organization", "logo": {{"@type": "ImageObject", "url": "https://pm.shiny.com.tw/images/logo.png"}}, "name": "Shiny黃金白銀"}}
}}
</script>'''

# 解析 docx 內容為 HTML 主體
def docx_to_html_body(docx_path):
    doc = docx.Document(docx_path)
    html = []
    headline = ""
    for para in doc.paragraphs:
        style = para.style.name.lower()
        text = para.text.strip()
        if not text:
            continue
        if "heading 1" in style:
            headline = text
            html.append(f"<h1>{text}</h1>")
        elif "heading 2" in style:
            html.append(f"<h2>{text}</h2>")
        elif "heading 3" in style:
            html.append(f"<h3>{text}</h3>")
        elif "list" in style:
            html.append(f"<ul><li>{text}</li></ul>")
        else:
            html.append(f"<p>{text}</p>")
    # 表格
    for table in doc.tables:
        html.append("<table>")
        for i, row in enumerate(table.rows):
            html.append("<tr>")
            for cell in row.cells:
                tag = "th" if i == 0 else "td"
                html.append(f"<{tag}>{cell.text.strip()}</{tag}>")
            html.append("</tr>")
        html.append("</table>")
    return "\n".join(html), headline

# GUI 主程式
class SEOHtmlApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word 轉 SEO HTML 工具")
        self.docx_path = ""
        # 欄位
        tk.Label(root, text="dateModified (YYYY-MM-DD)").grid(row=0, column=0)
        self.dateModified = tk.Entry(root)
        self.dateModified.grid(row=0, column=1)
        tk.Label(root, text="datePublished (YYYY-MM-DD)").grid(row=1, column=0)
        self.datePublished = tk.Entry(root)
        self.datePublished.grid(row=1, column=1)
        tk.Label(root, text="mainEntityOfPage id (數字)").grid(row=2, column=0)
        self.mainEntityId = tk.Entry(root)
        self.mainEntityId.grid(row=2, column=1)
        tk.Label(root, text="description").grid(row=3, column=0)
        self.description = tk.Entry(root, width=50)
        self.description.grid(row=3, column=1)
        # 選擇 Word 檔
        tk.Button(root, text="選擇 Word 檔", command=self.select_docx).grid(row=4, column=0)
        self.docx_label = tk.Label(root, text="尚未選擇檔案")
        self.docx_label.grid(row=4, column=1)
        # 產生 HTML
        tk.Button(root, text="產生 HTML", command=self.generate_html).grid(row=5, column=0, columnspan=2)

    def select_docx(self):
        path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if path:
            self.docx_path = path
            self.docx_label.config(text=os.path.basename(path))

    def generate_html(self):
        if not self.docx_path:
            messagebox.showerror("錯誤", "請先選擇 Word 檔")
            return
        dateModified = self.dateModified.get().strip()
        datePublished = self.datePublished.get().strip()
        mainEntityId = self.mainEntityId.get().strip()
        description = self.description.get().strip()
        html_body, headline = docx_to_html_body(self.docx_path)
        seo_script = SEO_TEMPLATE.format(
            dateModified=dateModified,
            datePublished=datePublished,
            mainEntityId=mainEntityId,
            description=description,
            headline=headline
        )
        html = seo_script + "\n<article class=\"seo-article-content\">\n" + html_body + "\n</article>"
        # 預設輸出檔名 output_YYYYMMDD[_N].html
        today_name = datetime.today().strftime('%Y%m%d')
        base = f"output_{today_name}.html"
        candidate = base
        idx = 1
        while os.path.exists(candidate):
            candidate = f"output_{today_name}_{idx}.html"
            idx += 1
        base_dir = os.path.dirname(os.path.dirname(__file__))
        out_path = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("HTML Files", "*.html")], initialfile=candidate, initialdir=os.path.join(base_dir, "output"))
        if out_path:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(html)
            messagebox.showinfo("完成", f"已產生 HTML：{out_path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = SEOHtmlApp(root)
    root.mainloop()
